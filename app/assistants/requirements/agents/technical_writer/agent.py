"""
技术文档编写师智能体 - 基于BaseAgent

专注于技术文档编写：
- 需求规格说明书
- 用户故事
- 验收标准
- 技术规范
"""
import os
from typing import Dict, List, Optional

from docx import Document
from weasyprint import HTML

from app.agent.base import BaseAgent
from app.llm import LLM
from app.logger import logger
from app.schema import ROLE_TYPE, Message


class TechnicalWriterAgent(BaseAgent):
    """技术文档编写师智能体"""

    def __init__(self, name: str = "技术文档编写师", **kwargs):
        system_prompt = self._get_system_prompt()

        super().__init__(
            name=name,
            system_prompt=system_prompt,
            next_step_prompt="请基于需求分析结果编写专业的技术文档。",
            llm=LLM(config_name="technical_writer"),
            max_steps=4,
            **kwargs,
        )

        self.document_sections: Dict[str, str] = {}
        self.user_stories: List[str] = []
        self.acceptance_criteria: List[str] = []

    def _get_system_prompt(self) -> str:
        """技术文档编写师系统提示词"""
        return """你是一位专业的技术文档编写师，专门负责编写高质量的软件需求文档。

# 职责
- 编写清晰、完整的需求规格说明书
- 创建用户故事和验收标准
- 制定技术规范和约束条件
- 确保文档的专业性和可读性

# 文档标准
1. **结构清晰**: 使用标准的文档结构
2. **内容完整**: 涵盖所有关键需求信息
3. **语言准确**: 使用准确、无歧义的技术语言
4. **可操作性**: 文档应便于开发团队理解和实施

# 文档模板
## 需求规格说明书结构
1. **项目概述**
2. **功能需求**
3. **非功能需求**
4. **系统约束**
5. **用户角色和权限**
6. **业务规则**
7. **验收标准**
8. **风险评估**

请基于需求分析结果，编写专业的技术文档。"""

    async def write_documentation(
        self, analysis_result: str, knowledge: Optional[Dict] = None
    ) -> str:
        """
        执行文档编写并返回结果

        Args:
            analysis_result: 分析结果
            knowledge: 知识库相关信息

        Returns:
            str: 文档内容
        """
        try:
            # 将分析结果添加到记忆
            self.update_memory(ROLE_TYPE.REQUIREMENT_CLARIFIER, f"业务分析结果：\n\n{analysis_result}")

            # 一次性生成完整文档
            return await self._compile_quick_document(analysis_result, knowledge)
        except Exception as e:
            logger.error(f"文档编写失败: {str(e)}")
            raise

    async def _compile_quick_document(
        self, analysis_result: str, knowledge: Optional[Dict] = None
    ) -> str:
        """快速生成完整文档"""
        prompt = f"""请基于以下业务分析结果，编写完整的需求规格说明书：

业务分析结果：
{analysis_result}

最终文档结构：
1. **项目概述**
2. **功能需求**
3. **非功能需求**
4. **系统约束和限制**
5. **验收标准和测试计划**
6. **项目里程碑和交付计划**
7. **风险评估和缓解措施**

请整合所有内容，生成完整、专业的需求规格说明书。
确保文档结构清晰、内容完整、便于开发团队使用。"""

        response = await self.llm.ask(
            [Message.system_message(prompt)], temperature=0.2
        )
        response_str = str(response)
        self.update_memory(ROLE_TYPE.TECHNICAL_WRITER, response_str)

        logger.info("需求规格说明书编写完成")
        return response_str

    async def step(self, content: Optional[str] = None) -> str:
        """执行文档编写步骤"""
        try:
            if self.current_step == 1:
                return await self._write_project_overview()
            elif self.current_step == 2:
                return await self._write_functional_requirements()
            elif self.current_step == 3:
                return await self._write_non_functional_requirements()
            else:
                return await self._compile_final_document()

        except Exception as e:
            logger.error(f"文档编写步骤执行失败: {e}")
            return f"文档编写过程中发生错误: {str(e)}"

    async def _write_project_overview(self) -> str:
        """编写项目概述"""
        prompt = """请基于需求分析结果编写项目概述部分：

内容包括：
1. **项目背景**: 项目的背景和起因
2. **项目目标**: 明确的项目目标和期望成果
3. **项目范围**: 包含和不包含的功能范围
4. **目标用户**: 主要用户群体和使用场景
5. **业务价值**: 项目预期带来的业务价值

请使用专业、清晰的技术文档写作风格。"""

        messages = [Message(**msg) for msg in self.memory.messages] + [Message.system_message(prompt)]

        response = await self.llm.ask(messages, temperature=0.2)
        response_str = str(response)
        self.update_memory(ROLE_TYPE.TECHNICAL_WRITER, response_str)
        self.document_sections["project_overview"] = response_str

        logger.info("完成项目概述编写")
        return response_str

    async def _write_functional_requirements(self) -> str:
        """编写功能需求"""
        prompt = """请编写详细的功能需求部分：

内容结构：
1. **核心功能模块**: 主要功能模块划分
2. **功能详细描述**: 每个功能的详细说明
3. **用户故事**: 以用户故事形式描述需求
4. **功能优先级**: 功能的重要性和优先级
5. **功能依赖关系**: 功能间的依赖关系

用户故事格式：
作为[用户角色]，我希望[功能描述]，以便[业务价值]。

请确保功能需求描述准确、完整、无歧义。"""

        messages = [Message(**msg) for msg in self.memory.messages] + [Message.system_message(prompt)]

        response = await self.llm.ask(messages, temperature=0.3)
        response_str = str(response)
        self.update_memory(ROLE_TYPE.TECHNICAL_WRITER, response_str)
        self.document_sections["functional_requirements"] = response_str

        logger.info("完成功能需求编写")
        return response_str

    async def _write_non_functional_requirements(self) -> str:
        """编写非功能需求"""
        prompt = """请编写非功能需求部分：

内容包括：
1. **性能需求**: 响应时间、吞吐量、并发用户数等
2. **可用性需求**: 系统可用性指标和要求
3. **安全需求**: 数据安全、访问控制、隐私保护
4. **兼容性需求**: 浏览器、操作系统、设备兼容性
5. **可维护性需求**: 代码质量、文档要求、测试覆盖率
6. **部署需求**: 部署环境、服务器配置、扩展性

请为每个非功能需求提供具体的指标和验收标准。"""

        messages = [Message(**msg) for msg in self.memory.messages] + [Message.system_message(prompt)]

        response = await self.llm.ask(messages, temperature=0.3)
        response_str = str(response)
        self.update_memory(ROLE_TYPE.TECHNICAL_WRITER, response_str)
        self.document_sections["non_functional_requirements"] = response_str

        logger.info("完成非功能需求编写")
        return response_str

    async def _compile_final_document(self) -> str:
        """编译最终文档"""
        prompt = """请基于前面编写的各个部分，编译完整的需求规格说明书：

最终文档结构：
1. **项目概述** (已完成)
2. **功能需求** (已完成)
3. **非功能需求** (已完成)
4. **系统约束和限制**
5. **验收标准和测试计划**
6. **项目里程碑和交付计划**
7. **风险评估和缓解措施**
8. **附录和参考资料**

请整合所有内容，生成完整、专业的需求规格说明书。
确保文档结构清晰、内容完整、便于开发团队使用。"""

        messages = [Message(**msg) for msg in self.memory.messages] + [Message.system_message(prompt)]

        response = await self.llm.ask(messages, temperature=0.2)
        response_str = str(response)
        self.update_memory(ROLE_TYPE.TECHNICAL_WRITER, response_str)

        # 导出文档
        await self.export_document(response_str, "docx")
        await self.export_document(response_str, "pdf")

        # 设置完成状态
        self.state = self.state.__class__.FINISHED

        logger.info("需求规格说明书编写完成")
        return response_str

    def get_document_summary(self) -> Dict:
        """获取文档摘要"""
        return {
            "sections_completed": list(self.document_sections.keys()),
            "user_stories_count": len(self.user_stories),
            "acceptance_criteria_count": len(self.acceptance_criteria),
            "current_step": self.current_step,
            "state": (
                self.state.value if hasattr(self.state, "value") else str(self.state)
            ),
        }

    def get_full_document(self) -> str:
        """获取完整文档"""
        if len(self.memory.messages) > 0:
            # 返回最后一条助手消息（完整文档）
            for message_dict in reversed(self.memory.messages):
                message = Message(**message_dict)
                if message.metadata and message.metadata.get('role') == "assistant" and len(message.content) > 1000:
                    return message.content
        return "文档尚未完成编写"

    async def export_document(self, content: str, format: str) -> str:
        """
        导出文档

        Args:
            content: 文档内容 (Markdown格式)
            format: 导出格式 (docx, pdf)

        Returns:
            str: 导出文件的路径
        """
        output_dir = "data/requirements_documents"
        os.makedirs(output_dir, exist_ok=True)
        timestamp = self.memory.timestamp
        file_path = os.path.join(output_dir, f"requirements_{timestamp}.{format}")

        if format == "docx":
            document = Document()
            document.add_heading("需求规格说明书", 0)
            # 此处需要一个Markdown到docx的转换器，或者手动解析Markdown并添加到文档中
            # 为简化起见，我们直接添加文本
            document.add_paragraph(content)
            document.save(file_path)
        elif format == "pdf":
            # WeasyPrint需要HTML格式的输入
            # 此处需要一个Markdown到HTML的转换器
            import markdown
            html = markdown.markdown(content)
            HTML(string=html).write_pdf(file_path)
        else:
            raise ValueError(f"不支持的导出格式: {format}")

        logger.info(f"文档已导出到: {file_path}")
        return file_path
