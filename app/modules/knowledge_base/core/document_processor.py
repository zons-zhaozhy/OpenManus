"""
文档处理器 - 支持多种文档类型的解析和知识点自动提取
"""

import asyncio
import hashlib
import json
import mimetypes
import os
import re
from dataclasses import asdict
from datetime import datetime
from pathlib import Path
from typing import Any, Dict, List, Optional, Tuple

from loguru import logger

# 导入文档解析库
try:
    import pdfplumber
    import PyPDF2

    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False
    logger.warning("PDF解析库未安装，将跳过PDF支持")

try:
    from docx import Document as DocxDocument

    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    logger.warning("DOCX解析库未安装，将跳过Word文档支持")

try:
    import pandas as pd

    EXCEL_AVAILABLE = True
except ImportError:
    EXCEL_AVAILABLE = False
    logger.warning("Excel解析库未安装，将跳过Excel支持")

try:
    from bs4 import BeautifulSoup

    HTML_AVAILABLE = True
except ImportError:
    HTML_AVAILABLE = False
    logger.warning("HTML解析库未安装，将跳过HTML支持")

try:
    import markdown

    MARKDOWN_AVAILABLE = True
except ImportError:
    MARKDOWN_AVAILABLE = False
    logger.warning("Markdown解析库未安装，将跳过Markdown支持")

from .knowledge_manager import DocumentType, KnowledgeDocument, KnowledgePoint


class DocumentProcessor:
    """文档处理器 - 自动解析和提取知识点"""

    def __init__(self, storage_path: str = "data/knowledge_bases"):
        """
        初始化文档处理器

        Args:
            storage_path: 存储路径
        """
        self.storage_path = Path(storage_path)
        self.documents_path = self.storage_path / "documents"
        self.uploads_path = self.storage_path / "uploads"

        # 支持的文档类型映射
        self.supported_types = {
            ".pdf": DocumentType.PDF,
            ".doc": DocumentType.WORD,
            ".docx": DocumentType.WORD,
            ".md": DocumentType.MARKDOWN,
            ".markdown": DocumentType.MARKDOWN,
            ".txt": DocumentType.TEXT,
            ".html": DocumentType.HTML,
            ".htm": DocumentType.HTML,
            ".xlsx": DocumentType.EXCEL,
            ".xls": DocumentType.EXCEL,
            ".pptx": DocumentType.POWERPOINT,
            ".ppt": DocumentType.POWERPOINT,
        }

        logger.info("文档处理器初始化完成")

    async def upload_document(
        self,
        file_path: str,
        knowledge_base_id: str,
        title: Optional[str] = None,
        metadata: Optional[Dict[str, Any]] = None,
    ) -> Optional[KnowledgeDocument]:
        """
        上传并处理文档

        Args:
            file_path: 文件路径
            knowledge_base_id: 知识库ID
            title: 文档标题（可选）
            metadata: 元数据（可选）

        Returns:
            Optional[KnowledgeDocument]: 处理后的文档对象
        """
        file_path = Path(file_path)

        if not file_path.exists():
            logger.error(f"文件不存在: {file_path}")
            return None

        # 检查文件类型
        file_ext = file_path.suffix.lower()
        if file_ext not in self.supported_types:
            logger.error(f"不支持的文件类型: {file_ext}")
            return None

        document_type = self.supported_types[file_ext]

        # 计算文件信息
        file_size = file_path.stat().st_size
        checksum = self._calculate_checksum(file_path)

        # 复制文件到uploads目录
        upload_filename = f"{knowledge_base_id}_{checksum}_{file_path.name}"
        upload_path = self.uploads_path / upload_filename

        import shutil

        shutil.copy2(file_path, upload_path)

        # 解析文档内容
        content = await self._parse_document(upload_path, document_type)
        if not content:
            logger.error(f"文档解析失败: {file_path}")
            upload_path.unlink()  # 删除上传的文件
            return None

        # 提取知识点
        knowledge_points = await self._extract_knowledge_points(content)

        # 提取关键词
        keywords = self._extract_keywords(content)

        # 生成摘要
        summary = self._generate_summary(content)

        # 创建文档对象
        document = KnowledgeDocument(
            id=f"{knowledge_base_id}_{checksum}",
            knowledge_base_id=knowledge_base_id,
            title=title or file_path.stem,
            content=content,
            document_type=document_type,
            file_path=str(upload_path),
            file_size=file_size,
            checksum=checksum,
            knowledge_points=[kp.content for kp in knowledge_points],
            keywords=keywords,
            summary=summary,
            created_at=datetime.now(),
            updated_at=datetime.now(),
            metadata=metadata or {},
        )

        # 保存文档
        self._save_document(document)

        # 保存知识点
        for kp in knowledge_points:
            self._save_knowledge_point(kp)

        logger.info(f"文档处理完成: {document.title} (ID: {document.id})")
        return document

    async def _parse_document(
        self, file_path: Path, document_type: DocumentType
    ) -> Optional[str]:
        """
        解析文档内容

        Args:
            file_path: 文件路径
            document_type: 文档类型

        Returns:
            Optional[str]: 解析出的文本内容
        """
        try:
            if document_type == DocumentType.PDF:
                return self._parse_pdf(file_path)
            elif document_type == DocumentType.WORD:
                return self._parse_word(file_path)
            elif document_type == DocumentType.MARKDOWN:
                return self._parse_markdown(file_path)
            elif document_type == DocumentType.TEXT:
                return self._parse_text(file_path)
            elif document_type == DocumentType.HTML:
                return self._parse_html(file_path)
            elif document_type == DocumentType.EXCEL:
                return self._parse_excel(file_path)
            else:
                logger.warning(f"暂不支持解析类型: {document_type}")
                return None
        except Exception as e:
            logger.error(f"文档解析异常: {file_path}, 错误: {e}")
            return None

    def _parse_pdf(self, file_path: Path) -> Optional[str]:
        """解析PDF文档"""
        if not PDF_AVAILABLE:
            logger.error("PDF解析库未安装")
            return None

        content = []

        # 尝试使用pdfplumber（更好的表格支持）
        try:
            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    text = page.extract_text()
                    if text:
                        content.append(text)
        except Exception as e:
            logger.warning(f"pdfplumber解析失败，尝试PyPDF2: {e}")

            # 回退到PyPDF2
            try:
                with open(file_path, "rb") as file:
                    pdf_reader = PyPDF2.PdfReader(file)
                    for page in pdf_reader.pages:
                        text = page.extract_text()
                        if text:
                            content.append(text)
            except Exception as e2:
                logger.error(f"PyPDF2解析也失败: {e2}")
                return None

        return "\n".join(content)

    def _parse_word(self, file_path: Path) -> Optional[str]:
        """解析Word文档"""
        if not DOCX_AVAILABLE:
            logger.error("DOCX解析库未安装")
            return None

        try:
            doc = DocxDocument(file_path)
            content = []

            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    content.append(paragraph.text.strip())

            # 处理表格
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        content.append(" | ".join(row_text))

            return "\n".join(content)
        except Exception as e:
            logger.error(f"Word文档解析失败: {e}")
            return None

    def _parse_markdown(self, file_path: Path) -> Optional[str]:
        """解析Markdown文档"""
        try:
            with open(file_path, "r", encoding="utf-8") as f:
                content = f.read()

            if MARKDOWN_AVAILABLE:
                # 转换为HTML然后提取纯文本
                html = markdown.markdown(content)
                if HTML_AVAILABLE:
                    soup = BeautifulSoup(html, "html.parser")
                    return soup.get_text()

            # 简单处理：移除Markdown标记
            content = re.sub(r"#{1,6}\s+", "", content)  # 标题
            content = re.sub(r"\*\*(.*?)\*\*", r"\1", content)  # 粗体
            content = re.sub(r"\*(.*?)\*", r"\1", content)  # 斜体
            content = re.sub(r"`(.*?)`", r"\1", content)  # 代码
            content = re.sub(r"\[([^\]]+)\]\([^)]+\)", r"\1", content)  # 链接

            return content
        except Exception as e:
            logger.error(f"Markdown解析失败: {e}")
            return None

    def _parse_text(self, file_path: Path) -> Optional[str]:
        """解析纯文本文档"""
        try:
            with open(file_path, "r", encoding="utf-8") as f:
                return f.read()
        except UnicodeDecodeError:
            # 尝试其他编码
            for encoding in ["gbk", "latin1", "cp1252"]:
                try:
                    with open(file_path, "r", encoding=encoding) as f:
                        return f.read()
                except UnicodeDecodeError:
                    continue
            logger.error(f"无法确定文本文件编码: {file_path}")
            return None
        except Exception as e:
            logger.error(f"文本解析失败: {e}")
            return None

    def _parse_html(self, file_path: Path) -> Optional[str]:
        """解析HTML文档"""
        if not HTML_AVAILABLE:
            logger.error("HTML解析库未安装")
            return None

        try:
            with open(file_path, "r", encoding="utf-8") as f:
                html_content = f.read()

            soup = BeautifulSoup(html_content, "html.parser")

            # 移除script和style标签
            for script in soup(["script", "style"]):
                script.decompose()

            return soup.get_text()
        except Exception as e:
            logger.error(f"HTML解析失败: {e}")
            return None

    def _parse_excel(self, file_path: Path) -> Optional[str]:
        """解析Excel文档"""
        if not EXCEL_AVAILABLE:
            logger.error("Excel解析库未安装")
            return None

        try:
            # 读取所有工作表
            excel_file = pd.ExcelFile(file_path)
            content = []

            for sheet_name in excel_file.sheet_names:
                df = pd.read_excel(file_path, sheet_name=sheet_name)

                # 添加工作表名
                content.append(f"工作表: {sheet_name}")

                # 转换为文本
                content.append(df.to_string(index=False))
                content.append("")  # 空行分隔

            return "\n".join(content)
        except Exception as e:
            logger.error(f"Excel解析失败: {e}")
            return None

    async def _extract_knowledge_points(self, content: str) -> List[KnowledgePoint]:
        """
        从文档内容中提取知识点

        Args:
            content: 文档内容

        Returns:
            List[KnowledgePoint]: 知识点列表
        """
        knowledge_points = []

        # 按段落分割
        paragraphs = [p.strip() for p in content.split("\n") if p.strip()]

        for i, paragraph in enumerate(paragraphs):
            # 跳过太短的段落
            if len(paragraph) < 20:
                continue

            # 计算重要性评分（简单规则）
            importance_score = self._calculate_importance_score(paragraph)

            # 只保留重要性较高的段落作为知识点
            if importance_score > 0.3:
                kp = KnowledgePoint(
                    id=f"kp_{hashlib.md5(paragraph.encode()).hexdigest()[:8]}",
                    document_id="",  # 稍后设置
                    content=paragraph,
                    category=self._classify_knowledge_point(paragraph),
                    importance_score=importance_score,
                    context=self._get_context(paragraphs, i),
                    related_points=[],
                    created_at=datetime.now(),
                )
                knowledge_points.append(kp)

        return knowledge_points

    def _calculate_importance_score(self, text: str) -> float:
        """
        计算文本重要性评分

        Args:
            text: 文本内容

        Returns:
            float: 重要性评分（0-1）
        """
        score = 0.0

        # 基础分数
        score += 0.2

        # 长度评分
        if 50 <= len(text) <= 500:
            score += 0.2
        elif len(text) > 500:
            score += 0.1

        # 关键词评分
        important_keywords = [
            "要求",
            "需求",
            "功能",
            "特性",
            "规格",
            "标准",
            "流程",
            "方法",
            "原则",
            "概念",
            "定义",
            "接口",
            "API",
            "数据",
            "结构",
            "架构",
            "设计",
            "实现",
            "测试",
            "部署",
            "配置",
            "安全",
            "性能",
            "优化",
        ]

        keyword_count = sum(1 for keyword in important_keywords if keyword in text)
        score += min(keyword_count * 0.1, 0.3)

        # 格式化文本评分（包含编号、标题等）
        if re.match(r"^\d+[\.\)]\s+", text) or re.match(
            r"^[一二三四五六七八九十]+[\.\)]\s+", text
        ):
            score += 0.1

        if text.endswith("：") or text.endswith(":"):
            score += 0.1

        return min(score, 1.0)

    def _classify_knowledge_point(self, text: str) -> str:
        """
        对知识点进行分类

        Args:
            text: 知识点文本

        Returns:
            str: 分类名称
        """
        categories = {
            "functional_requirement": ["功能", "特性", "业务", "操作", "处理"],
            "non_functional_requirement": [
                "性能",
                "安全",
                "可用性",
                "可靠性",
                "兼容性",
            ],
            "technical_specification": ["技术", "架构", "接口", "API", "协议"],
            "process_guideline": ["流程", "步骤", "方法", "过程", "规范"],
            "definition": ["定义", "概念", "术语", "含义"],
            "constraint": ["约束", "限制", "条件", "规则"],
        }

        for category, keywords in categories.items():
            if any(keyword in text for keyword in keywords):
                return category

        return "general"

    def _get_context(self, paragraphs: List[str], index: int, window: int = 2) -> str:
        """
        获取知识点的上下文

        Args:
            paragraphs: 段落列表
            index: 当前段落索引
            window: 上下文窗口大小

        Returns:
            str: 上下文文本
        """
        start = max(0, index - window)
        end = min(len(paragraphs), index + window + 1)

        context_paragraphs = paragraphs[start:end]
        return "\n".join(context_paragraphs)

    def _extract_keywords(self, content: str, max_keywords: int = 20) -> List[str]:
        """
        提取关键词

        Args:
            content: 文档内容
            max_keywords: 最大关键词数量

        Returns:
            List[str]: 关键词列表
        """
        # 简单的关键词提取（基于词频）
        words = re.findall(r"\b[\u4e00-\u9fff]+\b|\b[a-zA-Z]+\b", content)

        # 过滤停用词
        stop_words = {
            "的",
            "了",
            "在",
            "是",
            "我",
            "有",
            "和",
            "就",
            "不",
            "人",
            "都",
            "一",
            "一个",
            "上",
            "也",
            "很",
            "到",
            "说",
            "要",
            "去",
            "你",
            "会",
            "着",
            "没有",
            "看",
            "好",
            "自己",
            "这",
        }

        word_freq = {}
        for word in words:
            if len(word) > 1 and word not in stop_words:
                word_freq[word] = word_freq.get(word, 0) + 1

        # 按频率排序并返回前N个
        sorted_words = sorted(word_freq.items(), key=lambda x: x[1], reverse=True)
        return [word for word, freq in sorted_words[:max_keywords]]

    def _generate_summary(self, content: str, max_length: int = 200) -> str:
        """
        生成文档摘要

        Args:
            content: 文档内容
            max_length: 最大摘要长度

        Returns:
            str: 文档摘要
        """
        # 简单的摘要生成：取前几句话
        sentences = re.split(r"[。！？\.\!\?]", content)
        summary_sentences = []
        current_length = 0

        for sentence in sentences:
            sentence = sentence.strip()
            if not sentence:
                continue

            if current_length + len(sentence) <= max_length:
                summary_sentences.append(sentence)
                current_length += len(sentence)
            else:
                break

        summary = "。".join(summary_sentences)
        if summary and not summary.endswith("。"):
            summary += "。"

        return summary or content[:max_length] + "..."

    def _calculate_checksum(self, file_path: Path) -> str:
        """计算文件校验和"""
        hash_md5 = hashlib.md5()
        with open(file_path, "rb") as f:
            for chunk in iter(lambda: f.read(4096), b""):
                hash_md5.update(chunk)
        return hash_md5.hexdigest()

    def _save_document(self, document: KnowledgeDocument) -> None:
        """保存文档到存储"""
        doc_file = self.documents_path / f"{document.id}.json"

        # 转换为可序列化的字典
        data = asdict(document)
        data["document_type"] = document.document_type.value
        data["created_at"] = document.created_at.isoformat()
        data["updated_at"] = document.updated_at.isoformat()

        with open(doc_file, "w", encoding="utf-8") as f:
            json.dump(data, f, ensure_ascii=False, indent=2)

    def _save_knowledge_point(self, kp: KnowledgePoint) -> None:
        """保存知识点到存储"""
        kp_file = self.storage_path / "knowledge_points" / f"{kp.id}.json"
        kp_file.parent.mkdir(exist_ok=True)

        # 转换为可序列化的字典
        data = asdict(kp)
        data["created_at"] = kp.created_at.isoformat()

        with open(kp_file, "w", encoding="utf-8") as f:
            json.dump(data, f, ensure_ascii=False, indent=2)
